"""
Export module for Local File Deep Research.
"""

import os
import json
import logging
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Any, Union, BinaryIO
from datetime import datetime

# Configure logging
logger = logging.getLogger(__name__)

# Constants
EXPORTS_DIR = os.environ.get("EXPORTS_DIR", "exports")
# Use absolute path for exports directory
EXPORTS_DIR = os.path.abspath(EXPORTS_DIR)
logger.info(f"Export module using exports directory: {EXPORTS_DIR}")

class ExportError(Exception):
    """Base exception for export errors."""
    pass

def _ensure_exports_dir():
    """Ensure the exports directory exists."""
    try:
        logger.info(f"Ensuring exports directory exists: {EXPORTS_DIR}")
        os.makedirs(EXPORTS_DIR, exist_ok=True)

        # Check if directory is writable
        if not os.access(EXPORTS_DIR, os.W_OK):
            logger.warning(f"Exports directory {EXPORTS_DIR} is not writable")
            # Try to use a temporary directory as fallback
            global EXPORTS_DIR
            old_dir = EXPORTS_DIR
            EXPORTS_DIR = os.path.join(tempfile.gettempdir(), "exports")
            logger.info(f"Using temporary directory as fallback: {EXPORTS_DIR}")
            os.makedirs(EXPORTS_DIR, exist_ok=True)
            return

        logger.info(f"Exports directory ready: {EXPORTS_DIR}")
    except Exception as e:
        logger.error(f"Error ensuring exports directory: {e}")
        # Try to use a temporary directory as fallback
        global EXPORTS_DIR
        old_dir = EXPORTS_DIR
        EXPORTS_DIR = os.path.join(tempfile.gettempdir(), "exports")
        logger.info(f"Using temporary directory as fallback: {EXPORTS_DIR}")
        os.makedirs(EXPORTS_DIR, exist_ok=True)

def _get_export_path(export_id: str, format: str) -> str:
    """Get the path to an export file."""
    # Map format to file extension
    if format.lower() == "markdown":
        ext = "md"
    elif format.lower() == "html":
        ext = "html"
    elif format.lower() == "pdf":
        ext = "pdf"
    elif format.lower() == "docx":
        ext = "docx"
    else:
        ext = "txt"

    logger.info(f"Creating export path for format {format} with extension {ext}")
    return os.path.join(EXPORTS_DIR, f"{export_id}.{ext}")

def export_to_markdown(content: Dict[str, Any]) -> str:
    """
    Export content to Markdown format.

    Args:
        content: Content to export

    Returns:
        Markdown content
    """
    # Extract content
    title = content.get("title", "Untitled Report")
    report = content.get("report", "")
    sources = content.get("sources", [])
    findings = content.get("findings", [])

    # Build markdown
    markdown = f"# {title}\n\n"
    markdown += f"*Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n\n"

    # Add report
    markdown += report

    # Add sources
    if sources:
        markdown += "\n\n## Sources\n\n"
        for i, source in enumerate(sources, 1):
            if isinstance(source, dict):
                source_name = source.get("source_name", source.get("source_path", f"Source {i}"))
                markdown += f"{i}. {source_name}\n"
            else:
                markdown += f"{i}. {source}\n"

    # Add findings
    if findings:
        markdown += "\n\n## Detailed Findings\n\n"
        for i, finding in enumerate(findings, 1):
            markdown += f"### Finding {i}\n\n"

            # Add summary or content
            if "summary" in finding:
                markdown += f"**Summary:** {finding['summary']}\n\n"
            elif "content" in finding:
                markdown += f"**Content:** {finding['content'][:300]}...\n\n"

            # Add analysis if available
            if "analysis" in finding and isinstance(finding["analysis"], dict):
                analysis = finding["analysis"]

                if "key_points" in analysis and analysis["key_points"]:
                    markdown += "**Key Points:**\n\n"
                    for point in analysis["key_points"]:
                        markdown += f"- {point}\n"
                    markdown += "\n"

                if "entities" in analysis and analysis["entities"]:
                    markdown += "**Entities:**\n\n"
                    for entity in analysis["entities"]:
                        markdown += f"- {entity}\n"
                    markdown += "\n"

                if "sentiment" in analysis and analysis["sentiment"]:
                    markdown += f"**Sentiment:** {analysis['sentiment']}\n\n"

            # Add other fields
            for key in ["answer", "info", "cot", "fact_check"]:
                if key in finding:
                    markdown += f"**{key.title()}:** {finding[key]}\n\n"

            # Add citation
            if "citation" in finding:
                citation = finding["citation"]
                source_path = citation.get("source_path", "")
                source_name = citation.get("source_name", "")
                markdown += f"**Source:** {source_name or source_path}\n\n"

    return markdown

def export_to_html(content: Dict[str, Any]) -> str:
    """
    Export content to HTML format.

    Args:
        content: Content to export

    Returns:
        HTML content
    """
    # Convert to markdown first
    markdown = export_to_markdown(content)

    try:
        import markdown as md
        html = md.markdown(markdown, extensions=['tables', 'fenced_code'])
    except ImportError:
        # Simple markdown to HTML conversion if markdown package is not available
        html = markdown.replace("\n\n", "</p><p>")
        html = html.replace("# ", "<h1>").replace("\n## ", "</h1><h2>").replace("\n### ", "</h2><h3>")
        html = html.replace("**", "<strong>").replace("**", "</strong>")
        html = html.replace("*", "<em>").replace("*", "</em>")
        html = f"<p>{html}</p>"
        html = html.replace("<p><h1>", "<h1>").replace("</h1><p>", "</h1>")
        html = html.replace("<p><h2>", "<h2>").replace("</h2><p>", "</h2>")
        html = html.replace("<p><h3>", "<h3>").replace("</h3><p>", "</h3>")

    # Wrap in HTML document
    title = content.get("title", "Untitled Report")
    html_doc = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
        }}
        h1, h2, h3 {{
            color: #333;
        }}
        code {{
            background-color: #f5f5f5;
            padding: 2px 4px;
            border-radius: 4px;
        }}
        pre {{
            background-color: #f5f5f5;
            padding: 10px;
            border-radius: 4px;
            overflow-x: auto;
        }}
        blockquote {{
            border-left: 4px solid #ddd;
            padding-left: 10px;
            color: #666;
        }}
    </style>
</head>
<body>
    {html}
</body>
</html>
"""

    return html_doc

def export_to_pdf(content: Dict[str, Any]) -> bytes:
    """
    Export content to PDF format using PyFPDF.

    Args:
        content: Content to export

    Returns:
        PDF content as bytes

    Raises:
        ExportError: If PDF export fails
    """
    logger.info("Attempting to export to PDF format using PyFPDF")

    try:
        from fpdf import FPDF
        import re
        from io import BytesIO
        from datetime import datetime

        logger.info("Successfully imported PyFPDF modules")

        # Extract content
        title = content.get("title", "Untitled Report")
        report = content.get("report", "")
        sources = content.get("sources", [])
        findings = content.get("findings", [])

        # Create PDF object
        pdf = FPDF()
        pdf.add_page()

        # Set font
        pdf.set_font("Arial", "B", 16)

        # Add title
        pdf.cell(0, 10, title, 0, 1, "C")
        pdf.ln(5)

        # Add generation date
        pdf.set_font("Arial", "I", 10)
        pdf.cell(0, 10, f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", 0, 1, "C")
        pdf.ln(5)

        # Process markdown content
        pdf.set_font("Arial", "", 12)

        # Split content into paragraphs
        paragraphs = report.split('\n\n')
        for p in paragraphs:
            # Check if it's a heading
            if p.startswith('# '):
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, p[2:], 0, 1)
                pdf.set_font("Arial", "", 12)
            elif p.startswith('## '):
                pdf.set_font("Arial", "B", 13)
                pdf.cell(0, 10, p[3:], 0, 1)
                pdf.set_font("Arial", "", 12)
            elif p.startswith('### '):
                pdf.set_font("Arial", "B", 12)
                pdf.cell(0, 10, p[4:], 0, 1)
                pdf.set_font("Arial", "", 12)
            # Check if it's a list
            elif p.startswith('- ') or p.startswith('* '):
                lines = p.split('\n')
                for line in lines:
                    if line.startswith('- ') or line.startswith('* '):
                        pdf.cell(10, 10, "•", 0, 0)
                        pdf.multi_cell(0, 10, line[2:], 0, 1)
            # Regular paragraph
            else:
                # Clean up markdown formatting
                p = re.sub(r'\*\*(.*?)\*\*', r'\1', p)  # Bold
                p = re.sub(r'\*(.*?)\*', r'\1', p)  # Italic

                # Add paragraph with word wrap
                pdf.multi_cell(0, 10, p, 0, 1)
                pdf.ln(5)

        # Add sources if available
        if sources:
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, "Sources", 0, 1)
            pdf.set_font("Arial", "", 12)

            for i, source in enumerate(sources, 1):
                if isinstance(source, dict):
                    source_name = source.get("source_name", source.get("source_path", f"Source {i}"))
                    pdf.cell(0, 10, f"{i}. {source_name}", 0, 1)
                else:
                    pdf.cell(0, 10, f"{i}. {source}", 0, 1)

        # Add findings if available
        if findings:
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, "Detailed Findings", 0, 1)

            for i, finding in enumerate(findings, 1):
                pdf.set_font("Arial", "B", 12)
                pdf.cell(0, 10, f"Finding {i}", 0, 1)
                pdf.set_font("Arial", "", 12)

                # Add summary or content
                if "summary" in finding:
                    pdf.set_font("Arial", "B", 10)
                    pdf.cell(30, 10, "Summary:", 0, 0)
                    pdf.set_font("Arial", "", 10)
                    pdf.multi_cell(0, 10, finding["summary"], 0, 1)
                elif "content" in finding:
                    pdf.set_font("Arial", "B", 10)
                    pdf.cell(30, 10, "Content:", 0, 0)
                    pdf.set_font("Arial", "", 10)
                    pdf.multi_cell(0, 10, finding["content"][:300] + "...", 0, 1)

                # Add citation
                if "citation" in finding:
                    citation = finding["citation"]
                    source_path = citation.get("source_path", "")
                    source_name = citation.get("source_name", "")
                    pdf.set_font("Arial", "B", 10)
                    pdf.cell(30, 10, "Source:", 0, 0)
                    pdf.set_font("Arial", "", 10)
                    pdf.multi_cell(0, 10, source_name or source_path, 0, 1)

                pdf.ln(5)

        # Get PDF as bytes
        pdf_bytes = pdf.output(dest="S").encode('latin-1')

        logger.info(f"PDF export successful with PyFPDF, size: {len(pdf_bytes)} bytes")
        return pdf_bytes
    except ImportError as e:
        logger.error(f"PyFPDF not available: {e}")
        raise ExportError(f"PDF export requires PyFPDF to be installed: {e}")
    except Exception as e:
        logger.error(f"Error during PDF export: {e}")
        raise ExportError(f"Error during PDF export: {e}")

def export_to_docx(content: Dict[str, Any]) -> bytes:
    """
    Export content to DOCX format using python-docx.

    Args:
        content: Content to export

    Returns:
        DOCX content as bytes

    Raises:
        ExportError: If DOCX export fails
    """
    logger.info("Attempting to export to DOCX format using python-docx")

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        import re

        logger.info("Successfully imported python-docx modules")

        # Extract content
        title = content.get("title", "Untitled Report")
        report = content.get("report", "")
        sources = content.get("sources", [])
        findings = content.get("findings", [])

        logger.info(f"Creating DOCX document with title: {title}")

        # Create document
        doc = Document()
        logger.info("Document object created successfully")

        # Add title
        title_heading = doc.add_heading(title, 0)
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add generation date
        date_paragraph = doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add a separator
        doc.add_paragraph()

        # Process markdown in the report
        # Split the report into paragraphs
        paragraphs = report.split('\n\n')
        for p in paragraphs:
            # Check if it's a heading
            if p.startswith('# '):
                doc.add_heading(p[2:], 1)
            elif p.startswith('## '):
                doc.add_heading(p[3:], 2)
            elif p.startswith('### '):
                doc.add_heading(p[4:], 3)
            # Check if it's a list
            elif p.startswith('- ') or p.startswith('* '):
                lines = p.split('\n')
                for line in lines:
                    if line.startswith('- ') or line.startswith('* '):
                        doc.add_paragraph(line[2:], style='List Bullet')
            # Otherwise treat as normal paragraph
            else:
                # Create a new paragraph
                para = doc.add_paragraph()

                # Process markdown formatting
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', p)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        # Bold text
                        run = para.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        # Italic text
                        run = para.add_run(part[1:-1])
                        run.italic = True
                    elif part:
                        # Regular text
                        para.add_run(part)

        # Add sources if available
        if sources:
            doc.add_heading("Sources", 1)
            for i, source in enumerate(sources, 1):
                if isinstance(source, dict):
                    source_name = source.get("source_name", source.get("source_path", f"Source {i}"))
                    doc.add_paragraph(f"{i}. {source_name}")
                else:
                    doc.add_paragraph(f"{i}. {source}")

        # Add findings if available
        if findings:
            doc.add_heading("Detailed Findings", 1)
            for i, finding in enumerate(findings, 1):
                doc.add_heading(f"Finding {i}", 2)

                # Add summary or content
                if "summary" in finding:
                    p = doc.add_paragraph()
                    p.add_run("Summary: ").bold = True
                    p.add_run(finding["summary"])
                elif "content" in finding:
                    p = doc.add_paragraph()
                    p.add_run("Content: ").bold = True
                    p.add_run(f"{finding['content'][:300]}...")

                # Add citation
                if "citation" in finding:
                    citation = finding["citation"]
                    source_path = citation.get("source_path", "")
                    source_name = citation.get("source_name", "")
                    p = doc.add_paragraph()
                    p.add_run("Source: ").bold = True
                    p.add_run(source_name or source_path)

        logger.info("Document content added successfully, saving to bytes")

        # Save document to bytes
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        bytes_data = docx_bytes.read()
        logger.info(f"DOCX document saved successfully, size: {len(bytes_data)} bytes")
        return bytes_data
    except ImportError as e:
        logger.error(f"Failed to import python-docx: {e}")
        raise ExportError(f"DOCX export requires python-docx to be installed: {e}")
    except Exception as e:
        logger.error(f"Error during DOCX export: {e}")
        raise ExportError(f"Error during DOCX export: {e}")

def export_report(content: Dict[str, Any], format: str = "markdown") -> Union[str, bytes]:
    """
    Export a report to a specific format.

    Args:
        content: Report content
        format: Export format ("markdown", "html", "pdf", "docx")

    Returns:
        Exported content as string or bytes

    Raises:
        ExportError: If export fails
    """
    _ensure_exports_dir()

    if format == "markdown":
        return export_to_markdown(content)
    elif format == "html":
        return export_to_html(content)
    elif format == "pdf":
        return export_to_pdf(content)
    elif format == "docx":
        return export_to_docx(content)
    else:
        raise ExportError(f"Unsupported export format: {format}")

def save_export(content: Dict[str, Any], format: str = "markdown") -> str:
    """
    Export a report and save it to a file.

    Args:
        content: Report content
        format: Export format ("markdown", "html", "pdf", "docx")

    Returns:
        Path to the exported file

    Raises:
        ExportError: If export fails
    """
    logger.info(f"Starting save_export for format: {format}")
    logger.info(f"Content keys: {list(content.keys())}")

    _ensure_exports_dir()
    logger.info(f"Exports directory: {EXPORTS_DIR}")

    # Generate export ID
    title = content.get("title", "report")
    title = "".join(c if c.isalnum() else "_" for c in title).lower()
    export_id = f"{title}_{int(datetime.now().timestamp())}"
    logger.info(f"Generated export ID: {export_id}")

    # Get export path
    export_path = _get_export_path(export_id, format)
    logger.info(f"Export path: {export_path}")

    # Export content
    logger.info(f"Calling export_report for format: {format}")
    try:
        exported = export_report(content, format)
        if isinstance(exported, str):
            logger.info(f"Exported content as string, length: {len(exported)} characters")
        else:
            logger.info(f"Exported content as bytes, size: {len(exported)} bytes")
    except Exception as export_error:
        logger.error(f"Error in export_report: {export_error}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise ExportError(f"Failed to export content: {export_error}")

    # Save to file
    try:
        logger.info(f"Saving exported content to: {export_path}")
        if isinstance(exported, str):
            with open(export_path, "w", encoding="utf-8") as f:
                f.write(exported)
            logger.info(f"Successfully saved string content to file, size: {os.path.getsize(export_path)} bytes")
        else:
            with open(export_path, "wb") as f:
                f.write(exported)
            logger.info(f"Successfully saved binary content to file, size: {os.path.getsize(export_path)} bytes")

        # Verify the file was created
        if os.path.exists(export_path):
            logger.info(f"Verified file exists at: {export_path}")
        else:
            logger.error(f"File not found after save: {export_path}")
            raise ExportError(f"File not found after save: {export_path}")
    except Exception as e:
        logger.error(f"Failed to save export: {e}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise ExportError(f"Failed to save export: {e}")

    return export_path
