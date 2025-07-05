"""
Web UI for collaboration features in Local File Deep Research.
"""

import streamlit as st
import requests
import json
import os
import time
import io
import base64
import markdown
import logging
from datetime import datetime
import pandas as pd
import fitz  # PyMuPDF for better PDF processing

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Configuration
API_URL = os.environ.get("API_URL", "http://localhost:8006")

# Check if API is available
import requests
import time

def check_api_connection():
    """Check if the API is available."""
    try:
        response = requests.get(f"{API_URL}/health", timeout=2)
        if response.status_code == 200:
            return True
    except requests.exceptions.RequestException:
        pass
    return False

# Wait for API to start (max 10 seconds)
api_available = False
for _ in range(5):
    if check_api_connection():
        api_available = True
        break
    time.sleep(2)

if not api_available:
    import streamlit as st
    st.warning(f"Could not connect to API at {API_URL}. Make sure the API is running.")

# --- Helper Functions ---
def format_time(timestamp):
    """Format a timestamp into a readable date/time string."""
    if isinstance(timestamp, str):
        try:
            dt = datetime.fromisoformat(timestamp)
            return dt.strftime("%Y-%m-%d %H:%M:%S")
        except ValueError:
            return timestamp
    return timestamp

def get_auth_headers():
    """Get authentication headers for API requests."""
    if "auth_token" not in st.session_state:
        return {}
    return {"Authorization": f"Bearer {st.session_state.auth_token}"}

def api_request(method, endpoint, data=None, json=None, params=None, files=None):
    """Make an API request with authentication headers."""
    logger.info(f"API Request: {method.upper()} {endpoint}")

    if not api_available:
        error_msg = f"Cannot make API request to {endpoint} because the API is not available."
        logger.error(error_msg)
        st.error(error_msg)
        st.info("Make sure the API server is running at " + API_URL)
        return None

    headers = get_auth_headers()
    url = f"{API_URL}{endpoint}"
    logger.info(f"Full URL: {url}")

    # Log request details
    if json:
        logger.info(f"Request JSON: {json}")
    if data:
        logger.info(f"Request data: {data}")
    if params:
        logger.info(f"Request params: {params}")
    if files:
        logger.info(f"Request files: {list(files.keys())}")
    if headers:
        logger.info(f"Request headers: {headers}")

    try:
        if method.lower() == "get":
            logger.info(f"Sending GET request to {url}")
            # Use longer timeout for project document requests
            if endpoint.startswith("/projects/") and endpoint.endswith("/documents"):
                logger.info(f"Using longer timeout for project documents endpoint")
                response = requests.get(url, headers=headers, params=params, timeout=120)  # 2 minutes timeout for project documents
            elif endpoint.startswith("/projects/") and endpoint.endswith("/comments"):
                logger.info(f"Using longer timeout for project comments endpoint")
                response = requests.get(url, headers=headers, params=params, timeout=120)  # 2 minutes timeout for project comments
            else:
                response = requests.get(url, headers=headers, params=params, timeout=60)  # 1 minute for other GET requests
        elif method.lower() == "post":
            # Handle different types of data
            if files:
                # File uploads with form data
                logger.info(f"Sending POST request with files to {url}")
                response = requests.post(url, headers=headers, data=data, params=params, files=files, timeout=300)  # 5 minutes timeout for file uploads
            elif json is not None:
                # JSON data
                # Use a longer timeout for research requests
                if endpoint == "/research":
                    logger.info(f"Sending POST request with JSON to {url} (research endpoint)")
                    response = requests.post(url, headers=headers, json=json, params=params, timeout=600)  # 10 minutes timeout for research
                elif endpoint == "/export":
                    logger.info(f"Sending POST request with JSON to {url} (export endpoint)")
                    response = requests.post(url, headers=headers, json=json, params=params, timeout=300)  # 5 minutes timeout for export
                elif endpoint.startswith("/projects/") and endpoint.endswith("/documents"):
                    logger.info(f"Sending POST request with JSON to {url} (project documents endpoint)")
                    response = requests.post(url, headers=headers, json=json, params=params, timeout=300)  # 5 minutes timeout for project documents
                elif endpoint.startswith("/projects/") and endpoint.endswith("/comments"):
                    logger.info(f"Sending POST request with JSON to {url} (project comments endpoint)")
                    response = requests.post(url, headers=headers, json=json, params=params, timeout=300)  # 5 minutes timeout for project comments
                else:
                    logger.info(f"Sending POST request with JSON to {url}")
                    response = requests.post(url, headers=headers, json=json, params=params, timeout=120)  # 2 minutes for other requests
            else:
                # Form data
                logger.info(f"Sending POST request with form data to {url}")
                response = requests.post(url, headers=headers, data=data, params=params, timeout=30)
        elif method.lower() == "put":
            if json is not None:
                logger.info(f"Sending PUT request with JSON to {url}")
                response = requests.put(url, headers=headers, json=json, params=params, timeout=30)
            else:
                logger.info(f"Sending PUT request with form data to {url}")
                response = requests.put(url, headers=headers, data=data, params=params, timeout=30)
        elif method.lower() == "delete":
            logger.info(f"Sending DELETE request to {url}")
            response = requests.delete(url, headers=headers, params=params, timeout=30)
        else:
            error_msg = f"Unsupported method: {method}"
            logger.error(error_msg)
            st.error(error_msg)
            return None

        # Log response details
        logger.info(f"Response status code: {response.status_code}")
        logger.info(f"Response headers: {dict(response.headers)}")

        # Log the response for debugging
        if response.status_code >= 400:
            error_msg = f"API Error: {response.status_code} - {response.text}"
            logger.error(error_msg)
            st.error(error_msg)
            st.error(f"API request failed: {response.status_code} {response.reason} for url: {response.url}")
            return None

        response.raise_for_status()

        # Try to parse JSON response
        try:
            json_response = response.json()
            logger.info(f"Response JSON: {json_response}")
            return json_response
        except json.JSONDecodeError as json_error:
            error_msg = f"Failed to parse JSON response: {json_error}"
            logger.error(error_msg)
            logger.info(f"Response content: {response.text[:500]}...")
            st.error(error_msg)
            return None
    except requests.exceptions.ConnectionError as e:
        error_msg = f"Connection error: Could not connect to API at {API_URL}"
        logger.error(f"{error_msg}: {str(e)}")
        st.error(error_msg)
        st.info("Make sure the API server is running and accessible.")
        return None
    except requests.exceptions.Timeout as e:
        error_msg = f"Timeout error: The API request to {endpoint} timed out."
        logger.error(f"{error_msg}: {str(e)}")
        st.error(error_msg)
        return None
    except requests.exceptions.RequestException as e:
        error_msg = f"API request failed: {e}"
        logger.error(error_msg)
        st.error(error_msg)
        if hasattr(e, 'response') and e.response is not None and hasattr(e.response, 'text'):
            response_text = e.response.text
            logger.error(f"Response: {response_text}")
            st.error(f"Response: {response_text}")
        return None

# --- Authentication Functions ---
def login_form():
    """Display login form."""
    with st.form("login_form"):
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        submit = st.form_submit_button("Login")

        if submit and username and password:
            try:
                response = api_request("post", "/auth/login", json={"username": username, "password": password})
                if response and "token" in response:
                    st.session_state.auth_token = response["token"]
                    st.session_state.username = username
                    st.success("Login successful!")
                    st.rerun()
                else:
                    st.error("Login failed. Please check your credentials.")
            except Exception as e:
                st.error(f"Login failed: {str(e)}")
                st.info("Please check your credentials and try again.")

def register_form():
    """Display registration form."""
    with st.form("register_form"):
        username = st.text_input("Username")
        email = st.text_input("Email")
        password = st.text_input("Password", type="password")
        confirm_password = st.text_input("Confirm Password", type="password")
        submit = st.form_submit_button("Register")

        if submit:
            if not username or not email or not password:
                st.error("All fields are required.")
            elif password != confirm_password:
                st.error("Passwords do not match.")
            else:
                try:
                    response = api_request("post", "/auth/register", json={
                        "username": username,
                        "email": email,
                        "password": password,
                        "role": "user"
                    })
                    if response and "user" in response:
                        st.success("Registration successful! You can now login.")
                        st.session_state.show_login = True
                        st.rerun()
                    else:
                        st.error("Registration failed.")
                except Exception as e:
                    st.error(f"Registration failed: {str(e)}")
                    st.info("Please try again with a different username or email.")

def logout():
    """Logout the current user."""
    if "auth_token" in st.session_state:
        api_request("post", "/auth/logout")
        del st.session_state.auth_token
        del st.session_state.username
        st.success("Logged out successfully!")
        st.rerun()

def auth_page():
    """Display authentication page."""
    st.title("📚 Local File Deep Research")
    st.markdown("*Collaborative research platform for local files*")

    if "show_login" not in st.session_state:
        st.session_state.show_login = True

    tab1, tab2 = st.tabs(["Login", "Register"])

    with tab1:
        login_form()

    with tab2:
        register_form()

# --- Project Functions ---
def create_project_form():
    """Display form to create a new project."""
    with st.form("create_project_form"):
        name = st.text_input("Project Name")
        description = st.text_area("Description")
        submit = st.form_submit_button("Create Project")

        if submit and name:
            response = api_request("post", "/projects", json={
                "name": name,
                "description": description,
                "owner": st.session_state.username
            })
            if response and "project_id" in response:
                st.success(f"Project '{name}' created successfully!")
                st.session_state.current_project = response["project_id"]
                st.rerun()

def list_projects():
    """List projects for the current user."""
    projects = api_request("get", "/projects")
    if not projects:
        st.info("No projects found. Create a new project to get started.")
        return

    # Create a dataframe from projects
    project_data = []
    for project in projects:
        project_data.append({
            "Project Name": project["name"],
            "Description": project["description"],
            "Created": format_time(project["created_at"]),
            "Updated": format_time(project["updated_at"]),
            "Members": len(project["members"]),
            "Documents": len(project["documents"]),
            "Access": project["access_level"],
            "ID": project["project_id"]
        })

    if project_data:
        df = pd.DataFrame(project_data)
        st.dataframe(df, use_container_width=True, column_config={
            "Project Name": st.column_config.TextColumn("Project Name"),
            "Description": st.column_config.TextColumn("Description"),
            "Created": st.column_config.DatetimeColumn("Created"),
            "Updated": st.column_config.DatetimeColumn("Updated"),
            "Members": st.column_config.NumberColumn("Members"),
            "Documents": st.column_config.NumberColumn("Documents"),
            "Access": st.column_config.TextColumn("Access"),
            "ID": st.column_config.TextColumn("ID", required=False)
        })

        # Select project
        selected_project = st.selectbox("Select a project to view",
                                       [f"{p['Project Name']} ({p['ID']})" for p in project_data])
        if selected_project:
            project_id = selected_project.split("(")[-1].split(")")[0]
            st.session_state.current_project = project_id
            st.rerun()



def view_project(project_id):
    """View a specific project."""
    project = api_request("get", f"/projects/{project_id}")
    if not project:
        st.error(f"Project not found or you don't have access.")
        return

    st.header(project["name"])
    st.markdown(f"*{project['description']}*")
    st.markdown(f"**Created by:** {project['owner']} on {format_time(project['created_at'])}")
    st.markdown(f"**Last updated:** {format_time(project['updated_at'])}")

    # Add a prominent button to upload documents
    if st.button("📄 Upload New Document", use_container_width=True, type="primary"):
        st.session_state.show_upload_form = True

    # Show upload form if button was clicked
    if st.session_state.get("show_upload_form", False):
        with st.expander("Upload Document", expanded=True):
            uploaded_file = st.file_uploader("Upload a document", type=["pdf", "docx", "xlsx", "txt", "md", "py", "js", "html", "csv", "json", "pptx"], key="project_view_uploader")
            if uploaded_file:
                doc_title = st.text_input("Document Title", value=uploaded_file.name, key="project_view_doc_title")
                if st.button("Add Document", key="project_view_add_doc_btn"):
                    try:
                        # Use the unified file upload endpoint for all file types
                        with st.spinner(f"Processing {uploaded_file.name} using local file processing..."):
                            try:
                                # Get the bytes data from the uploaded file
                                bytes_data = uploaded_file.getvalue()

                                # Log the file size
                                st.info(f"File size: {len(bytes_data)} bytes")

                                # Prepare the file for upload
                                files = {"file": (uploaded_file.name, bytes_data, "application/octet-stream")}

                                # Send the request to the API with the file
                                response = api_request("post", f"/projects/{project_id}/file-upload",
                                                    data={"title": doc_title},
                                                    files=files)

                                # Display success message with file info
                                if response and "document_id" in response:
                                    db_path = response.get("db_file_path", "database")
                                    doc_type = response.get("document_type", "unknown")
                                    st.success(f"File '{doc_title}' ({doc_type}) processed and saved to {db_path}")
                                    st.info(f"Document ID: {response['document_id']}")
                            except Exception as e:
                                st.error(f"Error processing file: {str(e)}")
                                st.exception(e)
                                raise

                        if response and "document_id" in response:
                            st.success(f"Document '{doc_title}' added successfully!")
                            st.session_state.show_upload_form = False
                            st.rerun()
                    except Exception as e:
                        st.error(f"Error processing document: {str(e)}")
                        st.info("Try uploading a different file format or check if the file is corrupted.")

            # Add a button to cancel upload
            if st.button("Cancel Upload", key="cancel_upload_btn"):
                st.session_state.show_upload_form = False
                st.rerun()

    # Project tabs
    tab1, tab2, tab3, tab4 = st.tabs(["Documents", "Members", "Comments", "Shares"])

    # Documents tab
    with tab1:
        st.subheader("Project Documents")
        documents = api_request("get", f"/projects/{project_id}/documents")

        if not documents:
            st.info("No documents in this project yet.")
        else:
            doc_data = []
            for doc in documents:
                # Handle different document structures - some have title/document_type at the top level,
                # others have them nested in metadata
                title = doc.get("title", None)
                doc_type = doc.get("document_type", None)

                # If title or doc_type not found at top level, try to get from metadata
                if title is None and "metadata" in doc and isinstance(doc["metadata"], dict):
                    title = doc["metadata"].get("title", "Untitled")
                if doc_type is None and "metadata" in doc and isinstance(doc["metadata"], dict):
                    doc_type = doc["metadata"].get("document_type", "unknown")

                # Fallback values if still not found
                if title is None:
                    title = "Untitled"
                if doc_type is None:
                    doc_type = "unknown"

                doc_data.append({
                    "Title": title,
                    "Type": doc_type,
                    "Created": format_time(doc["created_at"]),
                    "Updated": format_time(doc["updated_at"]),
                    "Versions": doc["version_count"],
                    "ID": doc["document_id"]
                })

            if doc_data:
                st.dataframe(pd.DataFrame(doc_data), use_container_width=True)

                # Create two columns for document actions
                col1, col2 = st.columns(2)

                with col1:
                    # Select document to view
                    selected_doc = st.selectbox("Select a document to view",
                                              [f"{d['Title']} ({d['ID']})" for d in doc_data])
                    if selected_doc:
                        doc_id = selected_doc.split("(")[-1].split(")")[0]
                        st.session_state.current_document = doc_id
                        st.rerun()

                with col2:
                    # Select document to remove
                    selected_doc_to_remove = st.selectbox("Select a document to remove",
                                                        [f"{d['Title']} ({d['ID']})" for d in doc_data],
                                                        key="doc_to_remove")

                    # Add remove button with confirmation
                    if selected_doc_to_remove and st.button("🗑️ Remove Document", type="secondary"):
                        doc_id_to_remove = selected_doc_to_remove.split("(")[-1].split(")")[0]

                        # Show confirmation dialog
                        st.session_state.show_remove_confirmation = True
                        st.session_state.doc_id_to_remove = doc_id_to_remove
                        st.session_state.doc_title_to_remove = selected_doc_to_remove.split(" (")[0]
                        st.rerun()

                # Handle document removal confirmation
                if st.session_state.get("show_remove_confirmation", False):
                    doc_id = st.session_state.doc_id_to_remove
                    doc_title = st.session_state.doc_title_to_remove

                    st.warning(f"Are you sure you want to remove '{doc_title}' from this project?")

                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("Yes, Remove Document", key="confirm_remove"):
                            # Call API to remove document
                            try:
                                response = api_request("delete", f"/projects/{project_id}/documents/{doc_id}")
                                if response:
                                    st.success(f"Document '{doc_title}' removed successfully!")
                                    # Clear confirmation state
                                    st.session_state.show_remove_confirmation = False
                                    st.session_state.pop("doc_id_to_remove", None)
                                    st.session_state.pop("doc_title_to_remove", None)
                                    st.rerun()
                            except Exception as e:
                                st.error(f"Error removing document: {str(e)}")

                    with col2:
                        if st.button("Cancel", key="cancel_remove"):
                            # Clear confirmation state
                            st.session_state.show_remove_confirmation = False
                            st.session_state.pop("doc_id_to_remove", None)
                            st.session_state.pop("doc_title_to_remove", None)
                            st.rerun()

        # No need for a second upload section as we already have one at the top of the page

    # Members tab
    with tab2:
        st.subheader("Project Members")
        members = project["members"]

        if not members:
            st.info("No members in this project yet.")
        else:
            member_data = []
            for member in members:
                member_data.append({
                    "Username": member,
                    "Role": "Owner" if member == project["owner"] else "Member"
                })

            if member_data:
                st.dataframe(pd.DataFrame(member_data), use_container_width=True)

        # Add member
        if project["owner"] == st.session_state.username or project["access_level"] == "write":
            st.subheader("Add Member")
            new_member = st.text_input("Username")
            if st.button("Add Member") and new_member:
                response = api_request("post", f"/projects/{project_id}/members",
                                     json={"username": new_member})
                if response:
                    st.success(f"Member '{new_member}' added successfully!")
                    st.rerun()

    # Comments tab
    with tab3:
        st.subheader("Project Comments")
        comments = api_request("get", f"/projects/{project_id}/comments")

        if not comments:
            st.info("No comments in this project yet.")
        else:
            for comment in comments:
                with st.container():
                    st.markdown(f"**{comment['username']}** - {format_time(comment['created_at'])}")
                    st.markdown(comment["content"])
                    if comment["document_id"]:
                        st.markdown(f"*On document: {comment['document_id']}*")
                    st.divider()

        # Add comment
        st.subheader("Add Comment")
        comment_text = st.text_area("Comment")
        if st.button("Post Comment") and comment_text:
            response = api_request("post", f"/projects/{project_id}/comments",
                                 json={"content": comment_text})
            if response:
                st.success("Comment added successfully!")
                st.rerun()

    # Shares tab
    with tab4:
        st.subheader("Project Shares")
        shares = api_request("get", f"/projects/{project_id}/shares")

        if not shares:
            st.info("This project hasn't been shared yet.")
        else:
            share_data = []
            for share in shares:
                share_data.append({
                    "Shared With": share["recipient"],
                    "Shared By": share["sharer"],
                    "Permission": share["permission"],
                    "Created": format_time(share["created_at"]),
                    "ID": share["share_id"]
                })

            if share_data:
                st.dataframe(pd.DataFrame(share_data), use_container_width=True)

        # Share project
        if project["owner"] == st.session_state.username or project["access_level"] == "write":
            st.subheader("Share Project")
            recipient = st.text_input("Username to share with")
            permission = st.selectbox("Permission", ["read", "write"])
            if st.button("Share Project") and recipient:
                response = api_request("post", f"/projects/{project_id}/shares",
                                     json={"recipient": recipient, "permission": permission})
                if response:
                    st.success(f"Project shared with '{recipient}' successfully!")
                    st.rerun()

def upload_document_form():
    """Display form to upload a document."""
    uploaded_file = st.file_uploader(
        "Upload a document",
        type=["pdf", "docx", "xlsx", "txt", "md", "py", "js", "html", "csv", "json", "pptx"],
        key="global_document_uploader"
    )

    if uploaded_file:
        # Get projects for upload
        projects = api_request("get", "/projects")
        if not projects:
            st.error("You need to create a project first before uploading documents.")
            return

        # Create a dropdown for project selection
        upload_project_options = [(p["name"], p["project_id"]) for p in projects]
        selected_upload_project = st.selectbox(
            "Select project for upload",
            options=[name for name, _ in upload_project_options],
            key="global_upload_project_selector"
        )

        # Get the project ID for the selected project
        selected_upload_project_id = next((pid for name, pid in upload_project_options if name == selected_upload_project), None)

        doc_title = st.text_input("Document Title", value=uploaded_file.name, key="global_doc_title")

        if st.button("Add Document", key="global_add_doc_btn"):
            try:
                # Use the unified file upload endpoint for all file types
                with st.spinner(f"Processing {uploaded_file.name} using local file processing..."):
                    try:
                        # Get the bytes data from the uploaded file
                        bytes_data = uploaded_file.getvalue()

                        # Log the file size
                        st.info(f"File size: {len(bytes_data)} bytes")

                        # Prepare the file for upload
                        files = {"file": (uploaded_file.name, bytes_data, "application/octet-stream")}

                        # Send the request to the API with the file
                        response = api_request(
                            "post",
                            f"/projects/{selected_upload_project_id}/file-upload",
                            data={"title": doc_title},
                            files=files
                        )

                        # Display success message with file info
                        if response and "document_id" in response:
                            db_path = response.get("db_file_path", "database")
                            doc_type = response.get("document_type", "unknown")
                            st.success(f"File '{doc_title}' ({doc_type}) processed and saved to {db_path}")
                            st.info(f"Document ID: {response['document_id']}")
                    except Exception as e:
                        st.error(f"Error processing file: {str(e)}")
                        st.exception(e)
                        raise

                if response and "document_id" in response:
                    st.success(f"Document '{doc_title}' added successfully to project '{selected_upload_project}'!")
                    st.session_state.show_upload_form = False
                    st.rerun()
            except Exception as e:
                st.error(f"Error processing document: {str(e)}")
                st.info("Try uploading a different file format or check if the file is corrupted.")

    # Add a button to cancel upload
    if st.button("Cancel Upload", key="global_cancel_upload_btn"):
        st.session_state.show_upload_form = False
        st.rerun()

def view_document(document_id):
    """View a specific document."""
    document = api_request("get", f"/documents/{document_id}")
    if not document:
        st.error(f"Document not found or you don't have access.")
        return

    # Handle different document structures - some have title/document_type at the top level,
    # others have them nested in metadata
    title = document.get("title", None)
    doc_type = document.get("document_type", None)

    # If title or doc_type not found at top level, try to get from metadata
    if title is None and "metadata" in document and isinstance(document["metadata"], dict):
        title = document["metadata"].get("title", "Untitled")
    if doc_type is None and "metadata" in document and isinstance(document["metadata"], dict):
        doc_type = document["metadata"].get("document_type", "unknown")

    # Fallback values if still not found
    if title is None:
        title = "Untitled"
    if doc_type is None:
        doc_type = "unknown"

    # Add title and document_type to the document object for easier access later
    document["title"] = title
    document["document_type"] = doc_type

    st.header(title)

    # Add document actions
    _, col_actions2 = st.columns([3, 1])

    with col_actions2:
        # Add remove button
        if document.get("project_id") and st.button("🗑️ Remove from Project", type="secondary"):
            # Show confirmation dialog
            st.session_state.show_doc_view_remove_confirmation = True
            st.session_state.doc_view_id_to_remove = document_id
            st.session_state.doc_view_project_id = document.get("project_id")
            st.session_state.doc_view_title_to_remove = title
            st.rerun()

    # Handle document removal confirmation
    if st.session_state.get("show_doc_view_remove_confirmation", False):
        doc_id = st.session_state.doc_view_id_to_remove
        project_id = st.session_state.doc_view_project_id
        doc_title = st.session_state.doc_view_title_to_remove

        st.warning(f"Are you sure you want to remove '{doc_title}' from its project?")

        col1, col2 = st.columns(2)
        with col1:
            if st.button("Yes, Remove Document", key="confirm_doc_view_remove"):
                # Call API to remove document
                try:
                    response = api_request("delete", f"/projects/{project_id}/documents/{doc_id}")
                    if response:
                        st.success(f"Document '{doc_title}' removed successfully!")
                        # Clear confirmation state
                        st.session_state.show_doc_view_remove_confirmation = False
                        st.session_state.pop("doc_view_id_to_remove", None)
                        st.session_state.pop("doc_view_project_id", None)
                        st.session_state.pop("doc_view_title_to_remove", None)
                        # Go back to documents list
                        if "current_document" in st.session_state:
                            del st.session_state.current_document
                        st.rerun()
                except Exception as e:
                    st.error(f"Error removing document: {str(e)}")

        with col2:
            if st.button("Cancel", key="cancel_doc_view_remove"):
                # Clear confirmation state
                st.session_state.show_doc_view_remove_confirmation = False
                st.session_state.pop("doc_view_id_to_remove", None)
                st.session_state.pop("doc_view_project_id", None)
                st.session_state.pop("doc_view_title_to_remove", None)
                st.rerun()

    # Basic document info
    col1, col2 = st.columns(2)
    with col1:
        st.markdown(f"**Type:** {document['document_type']}")
        st.markdown(f"**Created:** {format_time(document['created_at'])}")
        st.markdown(f"**Last updated:** {format_time(document['updated_at'])}")
        st.markdown(f"**Owner:** {document.get('owner', 'Unknown')}")
        if document.get("project_id"):
            st.markdown(f"**Project ID:** {document.get('project_id')}")

    with col2:
        # Display file paths and metadata
        if 'file_path' in document:
            st.markdown(f"**File Path:** {document['file_path']}")
        if 'db_file_path' in document:
            st.markdown(f"**Database Path:** {document['db_file_path']}")
        if 'original_filename' in document:
            st.markdown(f"**Original Filename:** {document['original_filename']}")

    # Display metadata in an expander
    if 'metadata' in document and document['metadata']:
        with st.expander("Document Metadata"):
            for key, value in document['metadata'].items():
                st.markdown(f"**{key}:** {value}")

    # Document tabs
    tab1, tab2, tab3 = st.tabs(["Content", "Versions", "Comments"])

    # Content tab
    with tab1:
        st.subheader("Document Content")

        # Check if content is available
        if "content" in document:
            content = document["content"]

            # Display content based on document type
            if document["document_type"] in ["pdf", "docx", "pptx"]:
                st.markdown(content[:5000] + "..." if len(content) > 5000 else content)
            elif document["document_type"] in ["xlsx", "csv"]:
                try:
                    df = pd.read_csv(content) if document["document_type"] == "csv" else pd.read_excel(content)
                    st.dataframe(df, use_container_width=True)
                except:
                    st.markdown(content[:5000] + "..." if len(content) > 5000 else content)
            elif document["document_type"] in ["py", "js", "html", "java", "c", "cpp"]:
                st.code(content, language=document["document_type"])
            else:
                st.markdown(content[:5000] + "..." if len(content) > 5000 else content)
        elif "content_excerpt" in document:
            st.markdown("**Content Excerpt:**")
            st.markdown(document["content_excerpt"])
            st.info("Full content is not stored in the database to save space. Use the Research tab to search through the document content.")
        else:
            st.info("Content is not available in the database. This is expected as we've optimized storage by not storing full document content. Use the Research tab to search through the document content.")

            # Show file path if available
            if 'file_path' in document:
                st.markdown(f"**File Path:** {document['file_path']}")
            if 'db_file_path' in document:
                st.markdown(f"**Database Path:** {document['db_file_path']}")

            # Show metadata summary if available
            if 'metadata' in document and document['metadata']:
                st.markdown("**Document Information:**")
                if 'size_bytes' in document['metadata']:
                    st.markdown(f"- Size: {document['metadata']['size_bytes']} bytes")
                if 'paragraph_count' in document['metadata']:
                    st.markdown(f"- Paragraphs: {document['metadata']['paragraph_count']}")
                if 'word_count' in document:
                    st.markdown(f"- Words: {document['word_count']}")

        # Analyze document
        st.subheader("Document Analysis")
        if st.button("Analyze Document"):
            with st.spinner("Analyzing document..."):
                analysis = api_request("post", f"/documents/{document_id}/analyze")
                if analysis:
                    st.success("Analysis complete!")

                    # Display analysis results
                    if "summary" in analysis:
                        st.markdown(f"**Summary:** {analysis['summary']}")

                    if "key_points" in analysis:
                        st.markdown("**Key Points:**")
                        for point in analysis["key_points"]:
                            st.markdown(f"- {point}")

                    if "entities" in analysis:
                        st.markdown("**Entities:**")
                        for entity in analysis["entities"]:
                            st.markdown(f"- {entity}")

                    # Display other analysis fields based on document type
                    if document["document_type"] in ["py", "js", "java", "c", "cpp"] and "functions" in analysis:
                        st.markdown("**Functions:**")
                        for func in analysis["functions"]:
                            st.markdown(f"- {func}")

                    if "insights" in analysis:
                        st.markdown("**Insights:**")
                        st.markdown(analysis["insights"])

    # Versions tab
    with tab2:
        st.subheader("Document Versions")
        versions = api_request("get", f"/documents/{document_id}/versions")

        if not versions:
            st.info("No version history available.")
        else:
            version_data = []
            for version in versions:
                version_data.append({
                    "Version": version["version_id"],
                    "Created": format_time(version["created_at"]),
                    "Changes": version.get("changes_summary", ""),
                    "Created By": version.get("created_by", "Unknown")
                })

            if version_data:
                st.dataframe(pd.DataFrame(version_data), use_container_width=True)

                # Compare versions
                if len(version_data) > 1:
                    st.subheader("Compare Versions")
                    col1, col2 = st.columns(2)
                    with col1:
                        v1 = st.selectbox("Version 1", [v["Version"] for v in version_data])
                    with col2:
                        v2 = st.selectbox("Version 2", [v["Version"] for v in version_data if v["Version"] != v1])

                    if st.button("Compare"):
                        diff = api_request("get", f"/documents/{document_id}/diff",
                                         params={"version1": v1, "version2": v2})
                        if diff:
                            st.code("\n".join(diff["diff"]))

        # Upload new version
        st.subheader("Upload New Version")
        uploaded_file = st.file_uploader("Upload new version", type=[document["document_type"]])
        if uploaded_file:
            changes = st.text_area("Describe changes")
            if st.button("Upload Version"):
                try:
                    # Use local file processing for all file types
                    with st.spinner(f"Processing {uploaded_file.name} version using local file processing..."):
                        try:
                            # Get the bytes data from the uploaded file
                            bytes_data = uploaded_file.getvalue()

                            # Log the file size
                            st.info(f"Version file size: {len(bytes_data)} bytes")

                            # Prepare the file for upload
                            files = {"file": (uploaded_file.name, bytes_data, "application/octet-stream")}

                            # Send the request to the API with the file
                            response = api_request("post", f"/documents/{document_id}/versions",
                                                data={"changes_summary": changes},
                                                files=files)

                            # Display success message with file info
                            if response and "version_id" in response:
                                st.success(f"Version of '{document['title']}' processed and saved successfully!")
                                st.info(f"Version ID: {response.get('version_id', 'unknown')}")
                        except Exception as e:
                            st.error(f"Error processing version file: {str(e)}")
                            st.exception(e)
                            raise

                    if response and "version_id" in response:
                        st.success(f"New version uploaded successfully!")
                        st.rerun()
                except Exception as e:
                    st.error(f"Error processing document: {str(e)}")
                    st.info("Try uploading a different file format or check if the file is corrupted.")

    # Comments tab
    with tab3:
        st.subheader("Document Comments")
        comments = api_request("get", f"/documents/{document_id}/comments")

        if not comments:
            st.info("No comments on this document yet.")
        else:
            for comment in comments:
                with st.container():
                    st.markdown(f"**{comment['username']}** - {format_time(comment['created_at'])}")
                    st.markdown(comment["content"])
                    st.divider()

        # Add comment
        st.subheader("Add Comment")
        comment_text = st.text_area("Comment")
        if st.button("Post Comment") and comment_text:
            response = api_request("post", f"/documents/{document_id}/comments",
                                 json={"content": comment_text})
            if response:
                st.success("Comment added successfully!")
                st.rerun()

# --- Main UI ---
def main(set_config=False):
    """Main UI function.

    Args:
        set_config: Whether to set the page config. Set to False if called from another module
                   that has already set the page config.
    """
    # Page config is now set in the page_config.py module
    # This parameter is kept for backward compatibility
    if set_config:
        try:
            st.set_page_config(layout="wide", page_title="Local File Deep Research", page_icon="📚")
        except Exception as e:
            # Ignore errors if page config is already set
            pass

    # Check if user is authenticated
    if "auth_token" not in st.session_state:
        auth_page()
        return

    # Sidebar
    with st.sidebar:
        st.title("📚 Local File Research")
        st.markdown(f"*Logged in as: {st.session_state.username}*")

        # Navigation
        st.header("Navigation")
        nav_options = ["Projects", "Documents", "Research", "Settings"]
        nav_selection = st.radio("Go to", nav_options)

        # Add project selector in sidebar
        if nav_selection in ["Projects", "Documents", "Research"]:
            st.header("Quick Navigation")

            # Get all projects
            projects = api_request("get", "/projects")

            if projects and nav_selection in ["Projects", "Documents"]:
                # Create project selector for Projects and Documents pages
                project_options = [(p["name"], p["project_id"]) for p in projects]

                if nav_selection == "Projects":
                    st.subheader("Project Selection")
                    selected_project_name = st.selectbox(
                        "Select a project",
                        options=["All Projects"] + [name for name, _ in project_options],
                        key="sidebar_project_selector"
                    )

                    if selected_project_name != "All Projects":
                        selected_project_id = next((pid for name, pid in project_options if name == selected_project_name), None)
                        if selected_project_id:
                            # Set the current project in session state
                            if st.button("View Selected Project", key="view_selected_project_btn"):
                                st.session_state.current_project = selected_project_id
                                st.session_state.force_show_projects_list = False
                                st.rerun()

                # Add document selector for Documents page
                if nav_selection == "Documents":
                    # Get all documents
                    documents = api_request("get", "/documents")

                    if documents:
                        st.subheader("Document Selection")
                        # Create document options with search
                        doc_options = [(doc["title"], doc["document_id"]) for doc in documents]

                        # Add search box for documents
                        doc_search = st.text_input("Search documents", key="doc_search_input")

                        # Filter documents based on search
                        if doc_search:
                            filtered_docs = [(title, doc_id) for title, doc_id in doc_options
                                            if doc_search.lower() in title.lower()]
                        else:
                            filtered_docs = doc_options

                        # Show document selector
                        if filtered_docs:
                            selected_doc_title = st.selectbox(
                                "Select a document",
                                options=["All Documents"] + [title for title, _ in filtered_docs],
                                key="sidebar_doc_selector"
                            )

                            if selected_doc_title != "All Documents":
                                selected_doc_id = next((doc_id for title, doc_id in filtered_docs if title == selected_doc_title), None)
                                if selected_doc_id:
                                    # Set the current document in session state
                                    if st.button("View Selected Document", key="view_selected_doc_btn"):
                                        st.session_state.current_document = selected_doc_id
                                        st.session_state.force_show_documents_list = False
                                        st.rerun()

            # Add multi-project selector for Research page
            if nav_selection == "Research" and projects:
                st.subheader("Research Projects")

                # Create multi-select for projects
                project_names = [p["name"] for p in projects]
                selected_projects = st.multiselect(
                    "Select projects for research",
                    options=project_names,
                    default=[],
                    key="research_projects_multiselect"
                )

                # Store selected projects in session state
                if selected_projects:
                    st.session_state.selected_research_projects = selected_projects
                    st.info(f"Research will be performed on {len(selected_projects)} selected projects")
                else:
                    st.session_state.selected_research_projects = []
                    st.info("Research will be performed on all projects")

        # Logout button
        st.header("Account")
        if st.button("Logout", use_container_width=True):
            logout()

    # Initialize session state variables if they don't exist
    if "page" not in st.session_state:
        st.session_state.page = nav_selection

    # Initialize force_show flags if they don't exist
    if "force_show_projects_list" not in st.session_state:
        st.session_state.force_show_projects_list = True

    if "force_show_documents_list" not in st.session_state:
        st.session_state.force_show_documents_list = True

    # Check if navigation has changed
    if st.session_state.page != nav_selection:
        # Clear research results when navigating away from Research page
        if st.session_state.page == "Research" and nav_selection != "Research":
            if "research_results" in st.session_state:
                st.session_state.research_results = None

        # When navigating to Projects, force showing the projects list
        if nav_selection == "Projects":
            st.session_state.force_show_projects_list = True

        # When navigating to Documents, force showing the documents list
        if nav_selection == "Documents":
            st.session_state.force_show_documents_list = True

        # Update current page
        st.session_state.page = nav_selection

    # Main content based on navigation
    if nav_selection == "Projects":
        st.header("📂 Projects")

        # Create new project button
        if st.button("Create New Project", use_container_width=True):
            st.session_state.show_create_project = True

        # Show create project form
        if st.session_state.get("show_create_project", False):
            with st.expander("Create New Project", expanded=True):
                create_project_form()

        # Check if we should force showing the projects list
        if st.session_state.get("force_show_projects_list", False):
            # Reset the flag
            st.session_state.force_show_projects_list = False
            # Clear current project if set
            if "current_project" in st.session_state:
                del st.session_state.current_project

        # View current project or list projects
        if "current_project" in st.session_state:
            view_project(st.session_state.current_project)
            if st.button("Back to Projects List"):
                del st.session_state.current_project
                # Set flag to force showing projects list on next render
                st.session_state.force_show_projects_list = True
                st.rerun()
        else:
            list_projects()

            # Add a section for uploading documents directly from the projects list
            st.subheader("Quick Document Upload")
            st.info("Select a project above to view details or upload a document directly to a project.")

            # Get projects for upload
            projects = api_request("get", "/projects")
            if projects:
                # Create a dropdown for project selection
                upload_project_options = [(p["name"], p["project_id"]) for p in projects]
                selected_upload_project = st.selectbox(
                    "Select project for upload",
                    options=[name for name, _ in upload_project_options],
                    key="upload_project_selector"
                )

                # Get the project ID for the selected project
                selected_upload_project_id = next((pid for name, pid in upload_project_options if name == selected_upload_project), None)

                if selected_upload_project_id:
                    # File uploader
                    uploaded_file = st.file_uploader(
                        "Upload a document",
                        type=["pdf", "docx", "xlsx", "txt", "md", "py", "js", "html", "csv", "json", "pptx"],
                        key="projects_list_uploader"
                    )

                    if uploaded_file:
                        doc_title = st.text_input("Document Title", value=uploaded_file.name, key="projects_list_doc_title")
                        if st.button("Add Document to Project", key="projects_list_add_doc_btn"):
                            try:
                                # Use the unified file upload endpoint for all file types
                                with st.spinner(f"Processing {uploaded_file.name} using local file processing..."):
                                    try:
                                        # Get the bytes data from the uploaded file
                                        bytes_data = uploaded_file.getvalue()

                                        # Log the file size
                                        st.info(f"File size: {len(bytes_data)} bytes")

                                        # Prepare the file for upload
                                        files = {"file": (uploaded_file.name, bytes_data, "application/octet-stream")}

                                        # Send the request to the API with the file
                                        response = api_request("post", f"/projects/{selected_upload_project_id}/file-upload",
                                                            data={"title": doc_title},
                                                            files=files)

                                        # Display success message with file info
                                        if response and "document_id" in response:
                                            db_path = response.get("db_file_path", "database")
                                            doc_type = response.get("document_type", "unknown")
                                            st.success(f"File '{doc_title}' ({doc_type}) processed and saved to {db_path}")
                                            st.info(f"Document ID: {response['document_id']}")
                                    except Exception as e:
                                        st.error(f"Error processing file: {str(e)}")
                                        st.exception(e)
                                        raise

                                if response and "document_id" in response:
                                    st.success(f"Document '{doc_title}' added successfully to project '{selected_upload_project}'!")
                                    st.rerun()
                            except Exception as e:
                                st.error(f"Error processing document: {str(e)}")
                                st.info("Try uploading a different file format or check if the file is corrupted.")

    elif nav_selection == "Documents":
        st.header("📄 Documents")

        # Upload document button
        if st.button("Upload New Document", use_container_width=True):
            st.session_state.show_upload_form = True

        # Show upload form
        if st.session_state.get("show_upload_form", False):
            with st.expander("Upload Document", expanded=True):
                upload_document_form()

        # Check if we should force showing the documents list
        if st.session_state.get("force_show_documents_list", False):
            # Reset the flag
            st.session_state.force_show_documents_list = False
            # Clear current document if set
            if "current_document" in st.session_state:
                del st.session_state.current_document

        if "current_document" in st.session_state:
            view_document(st.session_state.current_document)
            if st.button("Back to Documents List"):
                del st.session_state.current_document
                # Set flag to force showing documents list on next render
                st.session_state.force_show_documents_list = True
                st.rerun()
        else:
            # List all documents across projects
            documents = api_request("get", "/documents")
            if not documents:
                st.info("No documents found.")
            else:
                doc_data = []
                for doc in documents:
                    doc_data.append({
                        "Title": doc["title"],
                        "Project": doc.get("project_name", "Unknown"),
                        "Type": doc["document_type"],
                        "Created": format_time(doc["created_at"]),
                        "Updated": format_time(doc["updated_at"]),
                        "ID": doc["document_id"]
                    })

                if doc_data:
                    st.dataframe(pd.DataFrame(doc_data), use_container_width=True)

                    # Create two columns for document actions
                    col1, col2 = st.columns(2)

                    with col1:
                        # Select document to view
                        selected_doc = st.selectbox("Select a document to view",
                                                  [f"{d['Title']} ({d['ID']})" for d in doc_data])
                        if selected_doc:
                            doc_id = selected_doc.split("(")[-1].split(")")[0]
                            st.session_state.current_document = doc_id
                            st.rerun()

                    with col2:
                        # Select document to remove from project
                        selected_doc_to_remove = st.selectbox("Select a document to remove",
                                                            [f"{d['Title']} ({d['ID']})" for d in doc_data],
                                                            key="doc_to_remove_all")

                        # Get project ID for the selected document
                        doc_id_to_remove = None
                        project_id_for_doc = None
                        doc_title_to_remove = None

                        if selected_doc_to_remove:
                            doc_id_to_remove = selected_doc_to_remove.split("(")[-1].split(")")[0]
                            doc_title_to_remove = selected_doc_to_remove.split(" (")[0]

                            # Find the project ID for this document
                            for doc in documents:
                                if doc["document_id"] == doc_id_to_remove:
                                    project_id_for_doc = doc.get("project_id")
                                    break

                        # Add remove button with confirmation
                        if doc_id_to_remove and project_id_for_doc and st.button("🗑️ Remove Document", type="secondary", key="remove_doc_all"):
                            # Show confirmation dialog
                            st.session_state.show_remove_confirmation_all = True
                            st.session_state.doc_id_to_remove_all = doc_id_to_remove
                            st.session_state.project_id_for_doc_all = project_id_for_doc
                            st.session_state.doc_title_to_remove_all = doc_title_to_remove
                            st.rerun()

                    # Handle document removal confirmation
                    if st.session_state.get("show_remove_confirmation_all", False):
                        doc_id = st.session_state.doc_id_to_remove_all
                        project_id = st.session_state.project_id_for_doc_all
                        doc_title = st.session_state.doc_title_to_remove_all

                        st.warning(f"Are you sure you want to remove '{doc_title}' from its project?")

                        col1, col2 = st.columns(2)
                        with col1:
                            if st.button("Yes, Remove Document", key="confirm_remove_all"):
                                # Call API to remove document
                                try:
                                    response = api_request("delete", f"/projects/{project_id}/documents/{doc_id}")
                                    if response:
                                        st.success(f"Document '{doc_title}' removed successfully!")
                                        # Clear confirmation state
                                        st.session_state.show_remove_confirmation_all = False
                                        st.session_state.pop("doc_id_to_remove_all", None)
                                        st.session_state.pop("project_id_for_doc_all", None)
                                        st.session_state.pop("doc_title_to_remove_all", None)
                                        st.rerun()
                                except Exception as e:
                                    st.error(f"Error removing document: {str(e)}")

                        with col2:
                            if st.button("Cancel", key="cancel_remove_all"):
                                # Clear confirmation state
                                st.session_state.show_remove_confirmation_all = False
                                st.session_state.pop("doc_id_to_remove_all", None)
                                st.session_state.pop("project_id_for_doc_all", None)
                                st.session_state.pop("doc_title_to_remove_all", None)
                                st.rerun()

    elif nav_selection == "Research":
        st.header("🔍 Research")

        # Research query
        query = st.text_area("Enter your research query:", height=100,
                           placeholder="e.g., 'Summarize the key points from all documents'")

        # Research options
        col1, col2, col3 = st.columns(3)
        with col1:
            research_mode = st.selectbox("Research Mode", ["rag", "multi_iteration"])
        with col2:
            report_mode = st.selectbox("Report Mode", ["normal", "chain_of_thought", "enhanced"])
        with col3:
            top_k = st.number_input("Results to Retrieve", min_value=1, max_value=50, value=10)

        # Show max_iterations option only when multi_iteration is selected
        if research_mode == "multi_iteration":
            max_iterations = st.number_input("Max Iterations", min_value=1, max_value=5, value=3,
                                           help="Maximum number of research iterations to perform")

        # Project filter - use the selected projects from sidebar if available
        if "selected_research_projects" in st.session_state and st.session_state.selected_research_projects:
            # Use the projects selected in the sidebar
            selected_projects = st.session_state.selected_research_projects
            st.success(f"Using {len(selected_projects)} projects selected in the sidebar: {', '.join(selected_projects)}")
            project_filter = selected_projects  # Pass the list of project names
        else:
            # Fallback to single project selection if sidebar selection is not used
            projects = api_request("get", "/projects")
            if projects:
                project_options = ["All Projects"] + [p["name"] for p in projects]
                selected_project = st.selectbox("Filter by Project", project_options)
                project_filter = None if selected_project == "All Projects" else selected_project

                # Show info about automatic indexing
                if selected_project != "All Projects":
                    st.info("Documents are automatically indexed when uploaded. You only need to manually index if you've added documents through other means or if you want to refresh the index.")
            else:
                project_filter = None

        # Index documents manually if needed
        if st.button("Index Documents", use_container_width=True):
            with st.spinner("Indexing documents... (this may take a few minutes for large documents)"):
                # Add a special case for document indexing with longer timeout
                try:
                    url = f"{API_URL}/documents/index"
                    headers = get_auth_headers()
                    params = {"project_id": project_filter}

                    logger.info(f"Sending POST request to {url} for document indexing with extended timeout")
                    response = requests.post(url, headers=headers, params=params, timeout=600)  # 10 minutes timeout for indexing

                    if response.status_code >= 400:
                        error_msg = f"API Error: {response.status_code} - {response.text}"
                        logger.error(error_msg)
                        st.error(error_msg)
                        index_result = None
                    else:
                        index_result = response.json()
                except Exception as e:
                    error_msg = f"Error during document indexing: {str(e)}"
                    logger.error(error_msg)
                    st.error(error_msg)
                    index_result = None

                if index_result:
                    st.success(f"Indexed {index_result.get('document_count', 0)} documents with {index_result.get('chunk_count', 0)} chunks.")
                    st.session_state.session_id = index_result.get("session_id")
                    st.info(f"Session ID: {st.session_state.session_id}")

        # Run research
        if st.button("Run Research", use_container_width=True, type="primary") and query:
            with st.spinner("Researching..."):
                # Prepare the request payload
                payload = {
                    "query": query,
                    "research_mode": research_mode,
                    "report_mode": report_mode,
                    "top_k": top_k,
                    "session_id": st.session_state.get("session_id")
                }

                # Handle project filter - use the selected projects from sidebar if available
                if "selected_research_projects" in st.session_state and st.session_state.selected_research_projects:
                    # Use multiple projects selected in the sidebar
                    payload["project_filter"] = st.session_state.selected_research_projects
                    st.info(f"Using {len(st.session_state.selected_research_projects)} projects for research: {', '.join(st.session_state.selected_research_projects)}")
                else:
                    # Use single project filter from the research page
                    payload["project_filter"] = project_filter

                # Add max_iterations if multi_iteration research mode is selected
                if research_mode == "multi_iteration" and 'max_iterations' in locals():
                    payload["max_iterations"] = max_iterations

                # Use a longer timeout for research requests
                try:
                    result = api_request("post", "/research", json=payload)
                except Exception as e:
                    st.error(f"Research failed: {str(e)}")
                    st.info("Try again with a simpler query or fewer documents.")
                    result = None

                if result:
                    st.success("Research complete!")

                    # Store the result in session state
                    st.session_state.research_results = result

                    # Clear any previous display before showing new results
                    st.rerun()

        # Display research results if they exist in session state
        if "research_results" in st.session_state and st.session_state.research_results:
            # Create a container for the results to ensure they're only displayed once
            results_container = st.container()
            with results_container:
                display_research_results(st.session_state.research_results)

            # Add a button to clear results
            if st.button("Clear Results", use_container_width=True):
                st.session_state.research_results = None
                st.rerun()

    elif nav_selection == "Settings":
        settings_page()

def display_research_results(result):
    """Display research results."""
    # Generate a unique timestamp for this display
    import time
    import uuid
    timestamp = int(time.time() * 1000)  # Milliseconds since epoch
    # Add a UUID to ensure uniqueness even if called multiple times in the same millisecond
    unique_suffix = str(uuid.uuid4())[:8]

    # Display research process information if available
    if "findings" in result and any("questions_by_iteration" in f for f in result.get("findings", [])):
        with st.expander("Research Process Details", expanded=True):
            st.subheader("🔍 Multi-Iteration Research Process")

            # Find the questions_by_iteration
            questions_by_iteration = None
            for finding in result.get("findings", []):
                if "questions_by_iteration" in finding:
                    questions_by_iteration = finding["questions_by_iteration"]
                    break

            if questions_by_iteration:
                st.write("This research was conducted using multiple iterations of questions to explore the topic in depth.")

                # Display iterations and questions
                for iteration, questions in questions_by_iteration.items():
                    st.markdown(f"**Iteration {int(iteration) + 1}:**")
                    for i, question in enumerate(questions, 1):
                        st.markdown(f"{i}. {question}")
                    st.markdown("")

                # Display performance metrics if available
                performance_metrics = None
                for finding in result.get("findings", []):
                    if "performance_metrics" in finding:
                        performance_metrics = finding["performance_metrics"]
                        break

                if performance_metrics:
                    st.subheader("⏱️ Performance Metrics")
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Total Time", f"{performance_metrics.get('total_time', 0):.2f}s")
                    with col2:
                        st.metric("Search Time", f"{performance_metrics.get('search_time', 0):.2f}s")
                    with col3:
                        st.metric("Analysis Time", f"{performance_metrics.get('analysis_time', 0):.2f}s")

    # Display report
    st.subheader("📝 Generated Report")
    st.markdown(result.get("report", "No report generated."))

    # Get the report ID for use in keys
    report_id = result.get('report_id', 'default')

    # Create a unique key for this specific instance
    unique_id = f"{report_id}_{timestamp}_{unique_suffix}"

    # Get the report content
    report_content = result.get("report", "No report content available.")
    report_title = result.get("title", f"Research Report {result.get('report_id', 'default')}")

    # Generate PDF and DOCX files immediately and store in session state
    if "pdf_bytes" not in st.session_state:
        st.session_state.pdf_bytes = {}
    if "docx_bytes" not in st.session_state:
        st.session_state.docx_bytes = {}

    # Generate PDF
    try:
        logger.info(f"Generating PDF for report")
        from fpdf import FPDF
        import re
        from datetime import datetime

        # Function to sanitize text for PyFPDF (replace Unicode characters)
        def sanitize_for_pdf(text):
            # Replace common Unicode characters with ASCII equivalents
            replacements = {
                '\u2019': "'",  # Right single quotation mark
                '\u2018': "'",  # Left single quotation mark
                '\u201c': '"',  # Left double quotation mark
                '\u201d': '"',  # Right double quotation mark
                '\u2013': '-',   # En dash
                '\u2014': '--',  # Em dash
                '\u2026': '...', # Ellipsis
                '\u00a0': ' ',   # Non-breaking space
                '\u00b0': 'deg', # Degree sign
                '\u00ae': '(R)', # Registered trademark
                '\u00a9': '(C)', # Copyright
                '\u00e9': 'e',   # e with acute
                '\u00e8': 'e',   # e with grave
                '\u00e0': 'a',   # a with grave
                '\u00e7': 'c',   # c with cedilla
                '\u00f1': 'n',   # n with tilde
            }

            for char, replacement in replacements.items():
                text = text.replace(char, replacement)

            # Replace any remaining non-Latin-1 characters with '?'
            return text.encode('latin-1', 'replace').decode('latin-1')

        # Sanitize title and content
        safe_title = sanitize_for_pdf(report_title)

        # Create PDF object
        pdf = FPDF()
        pdf.add_page()

        # Set font
        pdf.set_font("Arial", "B", 16)

        # Add title
        pdf.cell(0, 10, safe_title, 0, 1, "C")
        pdf.ln(5)

        # Add generation date
        pdf.set_font("Arial", "I", 10)
        pdf.cell(0, 10, f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", 0, 1, "C")
        pdf.ln(5)

        # Process markdown content
        pdf.set_font("Arial", "", 12)

        # Split content into paragraphs
        paragraphs = report_content.split('\n\n')
        for p in paragraphs:
            # Sanitize paragraph text
            p = sanitize_for_pdf(p)

            # Check if it's a heading
            if p.startswith('# '):
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, p[2:], 0, 1)
                pdf.set_font("Arial", "", 12)
            elif p.startswith('## '):
                pdf.set_font("Arial", "B", 13)
                pdf.cell(0, 10, p[3:], 0, 1)
                pdf.set_font("Arial", "", 12)
            elif p.startswith('### '):
                pdf.set_font("Arial", "B", 12)
                pdf.cell(0, 10, p[4:], 0, 1)
                pdf.set_font("Arial", "", 12)
            # Check if it's a list
            elif p.startswith('- ') or p.startswith('* '):
                lines = p.split('\n')
                for line in lines:
                    if line.startswith('- ') or line.startswith('* '):
                        pdf.cell(10, 10, "*", 0, 0)  # Using * instead of bullet for compatibility
                        pdf.multi_cell(0, 10, sanitize_for_pdf(line[2:]), 0, 1)
            # Regular paragraph
            else:
                # Clean up markdown formatting
                p = re.sub(r'\*\*(.*?)\*\*', r'\1', p)  # Bold
                p = re.sub(r'\*(.*?)\*', r'\1', p)  # Italic

                # Add paragraph with word wrap
                pdf.multi_cell(0, 10, p, 0, 1)
                pdf.ln(5)

        # Get PDF as bytes - use a different approach to handle Unicode
        try:
            # First try the standard approach
            pdf_bytes = pdf.output(dest="S").encode('latin-1')
        except UnicodeEncodeError:
            # If that fails, use a different approach with BytesIO
            import io
            buffer = io.BytesIO()
            pdf.output(buffer)
            pdf_bytes = buffer.getvalue()
            logger.info("Used BytesIO approach for PDF generation due to Unicode characters")
        logger.info(f"PDF generated successfully with PyFPDF, size: {len(pdf_bytes)} bytes")

        # Store in session state
        st.session_state.pdf_bytes[unique_id] = pdf_bytes
    except Exception as e:
        logger.error(f"Error generating PDF: {str(e)}")
        logger.error(f"PDF generation traceback: {traceback.format_exc()}")
        st.session_state.pdf_bytes[unique_id] = None

    # Generate DOCX
    try:
        logger.info(f"Generating DOCX for report")
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
        import re
        from datetime import datetime

        # Create a new document
        doc = Document()

        # Add title
        title = doc.add_heading(report_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add generation date
        date_paragraph = doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add a separator
        doc.add_paragraph()

        # Process markdown content
        paragraphs = report_content.split('\n\n')
        for p in paragraphs:
            # Check if it's a heading
            if p.startswith('# '):
                doc.add_heading(p[2:], 1)
            elif p.startswith('## '):
                doc.add_heading(p[3:], 2)
            elif p.startswith('### '):
                doc.add_heading(p[4:], 3)
            # Check if it's a list
            elif p.startswith('- ') or p.startswith('* '):
                lines = p.split('\n')
                for line in lines:
                    if line.startswith('- ') or line.startswith('* '):
                        doc.add_paragraph(line[2:], style='List Bullet')
            # Otherwise treat as normal paragraph
            else:
                # Create a new paragraph
                para = doc.add_paragraph()

                # Process markdown formatting
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', p)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        # Bold text
                        run = para.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        # Italic text
                        run = para.add_run(part[1:-1])
                        run.italic = True
                    elif part:
                        # Regular text
                        para.add_run(part)

        # Save to BytesIO
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        logger.info(f"DOCX generated successfully with python-docx")

        # Store in session state
        st.session_state.docx_bytes[unique_id] = docx_bytes.getvalue()
    except Exception as e:
        logger.error(f"Error generating DOCX: {str(e)}")
        logger.error(f"DOCX generation traceback: {traceback.format_exc()}")
        st.session_state.docx_bytes[unique_id] = None

    # Create tabs for Export Options and Sources/Findings
    main_tabs = st.tabs(["Export Options", "Sources & Findings"])

    # Tab 1: Export Options
    with main_tabs[0]:
        st.write("Export Options:")

        # Create direct download buttons for each format
        button_cols = st.columns(4)

        # Markdown Download - Direct and simple
        with button_cols[0]:
            st.download_button(
                label="📥 Markdown",
                data=report_content,
                file_name=f"report_{int(time.time())}.md",
                mime="text/markdown",
                key=f"md_download_{unique_id}"
            )

        # HTML Download - Direct
        with button_cols[1]:
            # Convert markdown to HTML
            html_content = markdown.markdown(report_content)
            html_doc = f"""<!DOCTYPE html>
            <html>
            <head>
                <title>{report_title}</title>
                <meta charset="utf-8">
                <style>
                    body {{ font-family: Arial, sans-serif; line-height: 1.6; max-width: 900px; margin: 0 auto; padding: 20px; }}
                    h1 {{ color: #2c3e50; }}
                    h2 {{ color: #3498db; border-bottom: 1px solid #eee; padding-bottom: 5px; }}
                    h3 {{ color: #2980b9; }}
                    pre {{ background-color: #f8f8f8; border: 1px solid #ddd; padding: 10px; border-radius: 5px; overflow-x: auto; }}
                    blockquote {{ border-left: 4px solid #ccc; padding-left: 15px; color: #555; }}
                    table {{ border-collapse: collapse; width: 100%; }}
                    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                    th {{ background-color: #f2f2f2; }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>"""

            st.download_button(
                label="📥 HTML",
                data=html_doc,
                file_name=f"report_{int(time.time())}.html",
                mime="text/html",
                key=f"html_download_{unique_id}"
            )

        # PDF Download - Direct download button for pre-generated PDF
        with button_cols[2]:
            # Check if PDF was generated successfully
            if unique_id in st.session_state.pdf_bytes and st.session_state.pdf_bytes[unique_id] is not None:
                st.download_button(
                    label="📥 PDF",
                    data=st.session_state.pdf_bytes[unique_id],
                    file_name=f"report_{int(time.time())}.pdf",
                    mime="application/pdf",
                    key=f"pdf_download_{unique_id}"
                )
                logger.info(f"PDF download button created for report")
            else:
                st.error("PDF generation failed. Please try again.")

        # DOCX Download - Direct download button for pre-generated DOCX
        with button_cols[3]:
            # Check if DOCX was generated successfully
            if unique_id in st.session_state.docx_bytes and st.session_state.docx_bytes[unique_id] is not None:
                st.download_button(
                    label="📥 DOCX",
                    data=st.session_state.docx_bytes[unique_id],
                    file_name=f"report_{int(time.time())}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"docx_download_{unique_id}"
                )
                logger.info(f"DOCX download button created for report")
            else:
                st.error("DOCX generation failed. Please try again.")

    # Tab 2: Sources and Findings
    with main_tabs[1]:
        # Display sources and findings
        sources_findings_tabs = st.tabs(["Sources", "Detailed Findings"])
        with sources_findings_tabs[0]:
            sources = result.get("sources", [])
            if sources:
                for i, src in enumerate(sources, 1):
                    st.markdown(f"**[{i}]** {src}")
            else:
                st.info("No sources cited in this report.")

        with sources_findings_tabs[1]:
            findings = result.get("findings", [])
            if findings:
                for i, finding in enumerate(findings, 1):
                    with st.expander(f"Finding {i}"):
                        if "summary" in finding:
                            st.markdown(f"**Summary:** {finding['summary']}")
                        if "content" in finding:
                            st.markdown(f"**Content:** {finding['content'][:300]}...")
                        if "source" in finding:
                            st.markdown(f"**Source:** {finding['source']}")
            else:
                st.info("No detailed findings available.")

    # Note: Export Options tab is now the first tab (main_tabs[0])
    # No need for duplicate export options here


    # Debug buttons moved to the bottom of the page

    # Add a debug button to show the raw result
    if st.button("Show Raw Result", key=f"show_raw_{result.get('report_id', 'default')}_{timestamp}"):
        st.json(result)

    # Add a test export button with format selection
    if st.button("Test Export API", key=f"test_export_{result.get('report_id', 'default')}_{timestamp}"):
        st.info("Testing export API with a test report...")
        try:
            # Check if exports directory exists
            import os
            if not os.path.exists("exports"):
                os.makedirs("exports")
                logger.info("Created exports directory")
                st.info("Created exports directory")
            else:
                logger.info("Exports directory exists")
                st.info(f"Exports directory exists at {os.path.abspath('exports')}")

                # List files in exports directory
                files = os.listdir("exports")
                logger.info(f"Files in exports directory: {files}")
                st.info(f"Files in exports directory: {files}")

            # Use PDF as the test format
            test_format = "pdf"

            # Call the API to export a test report
            export_data = {
                "report_id": "test",
                "format": test_format,
                "content": "# Test Report\n\nThis is a test report generated by the Test Export API button.\n\n## Section 1\n\nThis is section 1 content.\n\n## Section 2\n\nThis is section 2 content.",
                "title": "Test Export Report"
            }
            logger.info(f"Test export request data: {export_data}")
            st.info(f"Sending test export request to {API_URL}/export with format: {test_format}")

            # Make the API request with explicit timeout
            try:
                export_result = api_request("post", "/export", json=export_data)
                logger.info(f"Test export API response: {export_result}")
            except Exception as api_error:
                logger.error(f"API request failed: {str(api_error)}")
                st.error(f"API request failed: {str(api_error)}")
                import traceback
                logger.error(f"API request traceback: {traceback.format_exc()}")
                return

            # Display the raw API response for debugging
            st.json(export_result)

            if export_result and "download_url" in export_result:
                # Get the download URL
                download_url = f"{API_URL}{export_result['download_url']}"
                logger.info(f"Generated test download URL: {download_url}")

                # Check if the file exists
                file_path = export_result.get("file_path", "")
                if os.path.exists(file_path):
                    file_size = os.path.getsize(file_path)
                    logger.info(f"Export file exists at {file_path}, size: {file_size} bytes")
                    st.success(f"Export file exists at {file_path}, size: {file_size} bytes")
                else:
                    logger.error(f"Export file does not exist at {file_path}")
                    st.error(f"Export file does not exist at {file_path}")

                # Create a direct download link
                st.success(f"Test {test_format.upper()} generated successfully!")
                st.markdown(f"""<a href="{download_url}" target="_blank" download>Click here to download your test {test_format.upper()} report</a>""", unsafe_allow_html=True)

                # Add a direct link to open in browser
                st.markdown(f"""<a href="{download_url}" target="_blank">Open in browser</a>""", unsafe_allow_html=True)
            else:
                st.error(f"Failed to generate test {test_format.upper()}: {export_result.get('detail', 'Unknown error')}")
        except Exception as e:
            st.error(f"Error testing export API: {str(e)}")
            logger.error(f"Error testing export API: {str(e)}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")



def settings_page():
    """Display settings page."""
    st.header("⚙️ Settings")

    # User profile
    st.subheader("User Profile")
    user_info = api_request("get", "/auth/me")
    if user_info:
        st.markdown(f"**Username:** {user_info['username']}")
        st.markdown(f"**Email:** {user_info['email']}")
        st.markdown(f"**Role:** {user_info['role']}")

    # Change password
    st.subheader("Change Password")
    with st.form("change_password_form"):
        current_password = st.text_input("Current Password", type="password")
        new_password = st.text_input("New Password", type="password")
        confirm_password = st.text_input("Confirm New Password", type="password")
        submit = st.form_submit_button("Change Password")

        if submit:
            if not current_password or not new_password:
                st.error("All fields are required.")
            elif new_password != confirm_password:
                st.error("New passwords do not match.")
            else:
                response = api_request("post", "/auth/change-password", data={
                    "current_password": current_password,
                    "new_password": new_password
                })
                if response and "message" in response:
                    st.success("Password changed successfully!")

    # Application settings
    st.subheader("Application Settings")

    # Theme
    theme = st.selectbox("Theme", ["Light", "Dark", "System"])
    if st.button("Apply Theme"):
        # This would need to be implemented with custom CSS or using Streamlit's theming
        st.success(f"{theme} theme applied!")

    # DSPy settings
    st.subheader("DSPy Settings")
    llm_provider = st.selectbox("LLM Provider", ["openai", "ollama", "anthropic"])
    llm_model = st.text_input("LLM Model", value="gpt-3.5-turbo" if llm_provider == "openai" else "llama3" if llm_provider == "ollama" else "claude-3-sonnet-20240229")
    api_base = st.text_input("API Base URL (optional)")

    if st.button("Save DSPy Settings"):
        response = api_request("post", "/settings/dspy", data={
            "llm_provider": llm_provider,
            "llm_model": llm_model,
            "api_base": api_base
        })
        if response:
            st.success("DSPy settings saved successfully!")

if __name__ == "__main__":
    main()
